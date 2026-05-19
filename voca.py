import streamlit as st
from google import genai
from google.genai import types
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import io
import time
import threading

# ==========================================
# 1. 워드 파일 디자인 & 생성 헬퍼 함수
# ==========================================
def set_cell_borders(cell, color="D9D9D9", sz="4", val="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
    tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def create_word_document(all_word_data):
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10.5)
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("통합 본문 단어장")
    title_run.font.bold = True
    title_run.font.size = Pt(18)
    title_p.paragraph_format.space_after = Pt(24)
    
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    col_widths = [Inches(1.5), Inches(2.0), Inches(4.0)]
    
    hdr_cells = table.rows[0].cells
    headers = ["본문 단어", "우리말 뜻", "영영 풀이"]
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        hdr_cells[i].width = col_widths[i]
        set_cell_shading(hdr_cells[i], "4F81BD")
        set_cell_borders(hdr_cells[i], color="A6A6A6")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = docx.shared.RGBColor(255, 255, 255)
            
    for row_idx, item in enumerate(all_word_data):
        row_cells = table.add_row().cells
        row_data = [item.get("word", ""), item.get("meaning", ""), item.get("definition", "")]
        
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            row_cells[i].width = col_widths[i]
            if row_idx % 2 == 1:
                set_cell_shading(row_cells[i], "F2F5F8")
            set_cell_borders(row_cells[i], color="D9D9D9")
            
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(10)
                
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 2. 백그라운드 AI 호출을 위한 스레드 워커 함수
# ==========================================
def gemini_api_worker(client, image_bytes, mime_type, prompt, result_container):
    """실제 AI 연산을 백그라운드에서 실행하여 메인 UI 스레드 멈춤을 방지합니다."""
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                types.Part.from_bytes(data=image_bytes, mime_type=mime_type),
                prompt
            ],
            config=types.GenerateContentConfig(response_mime_type="application/json")
        )
        result_container["data"] = json.loads(response.text)
        result_container["status"] = "success"
    except Exception as e:
        result_container["status"] = "error"
        result_container["error_msg"] = str(e)

# ==========================================
# 3. 이미지 비동기 분석 및 균일 속도 트래픽 처리 로직
# ==========================================
def process_images_safely(client, uploaded_files, api_key, progress_bar, status_text):
    all_data = []
    total_files = len(uploaded_files)
    
    prompt = """
    이 이미지에서 영어 단어, 우리말 뜻, 영영 풀이를 추출해서 정확한 JSON 배열 형식으로 출력해줘.
    필기구로 수정한 흔적이나 추가로 적은 필기는 무시하고, 원래 인쇄되어 있던 텍스트만 추출해줘.
    결과는 오직 아래 구조를 가진 JSON 데이터만 반환해야 해:
    [
      {"word": "단어", "meaning": "품사 및 뜻", "definition": "영영 풀이 내용"}
    ]
    """
    
    # 육안으로 보기에 완벽히 평탄한 전진 속도를 연출하기 위한 루프 제어
    ui_progress = 0.0
    
    for idx, file in enumerate(uploaded_files):
        file.seek(0)
        image_bytes = file.read()
        
        # 결과를 받아올 안전 금고(컨테이너) 생성
        worker_result = {"status": "pending", "data": None, "error_msg": None}
        
        # 백그라운드 비동기 스레드 생성 및 즉시 시작
        api_thread = threading.Thread(
            target=gemini_api_worker, 
            args=(client, image_bytes, file.type, prompt, worker_result)
        )
        api_thread.start()
        
        # 이 파일이 도달해야 하는 최종 목적지 상한선
        target_max_progress = (idx + 1) / total_files
        
        # AI가 일하는 동안 초당 일정한 속도로 부드럽게 게이지를 밀어줍니다.
        # 루프 한 번당 0.05초 대기 -> 1초에 약 20번 갱신
        while api_thread.is_alive():
            # 목표치 직전(95% 수준)까지는 일정한 등속도로 전진
            available_room = target_max_progress - ui_progress
            if available_room > (1.0 / total_files) * 0.15:
                # 파일당 할당량에 맞추어 매끄러운 속도로 증가
                ui_progress += (1.0 / total_files) * 0.006
            else:
                # 거의 다 왔는데 아직 AI 응답이 안 왔다면 미세하게 숨고르기 전진 (멈춘 느낌 차단)
                ui_progress += (1.0 / total_files) * 0.0008
                
            if ui_progress > 0.99: ui_progress = 0.99
            
            progress_bar.progress(ui_progress)
            status_text.markdown(f"🧠 **[ {int(ui_progress * 100)}% / 100% ]** ({idx+1}/{total_files}장) 인공지능이 영단어 매핑을 실시간 연산 중입니다..")
            time.sleep(0.05)
            
        # 스레드가 종료된 후 결과 데이터 처리
        if worker_result["status"] == "success" and worker_result["data"]:
            all_data.extend(worker_result["data"])
            
            # AI 처리가 끝난 후 다음 파일로 넘어가기 전 해당 파일의 목표 진행률로 부드럽게 보정 안착
            while ui_progress < target_max_progress:
                ui_progress += 0.01
                if ui_progress > target_max_progress: ui_progress = target_max_progress
                progress_bar.progress(ui_progress)
                status_text.markdown(f"📝 **[ {int(ui_progress * 100)}% / 100% ]** ({idx+1}/{total_files}장) 단어 통합 정제 완료!")
                time.sleep(0.01)
                
        elif worker_result["status"] == "error":
            error_msg = worker_result["error_msg"]
            if "quota" in error_msg.lower() or "limit" in error_msg.lower():
                if idx > 0:
                    st.warning("⚠️ 구글 계정의 하루 무료 사용량(20장)이 모두 마감되었습니다. 프로그램 보호를 위해 현재까지 변환된 파일들로만 워드를 생성합니다.")
                    break
                else:
                    st.error("❌ 오늘 사용 가능한 구글 무료 제공량(20장)을 모두 초과하여 변환을 시작할 수 없습니다. 내일 다시 시도해 주세요.")
                    return None
            else:
                st.error(f"❌ 변환 중 오류가 발생했습니다: {error_msg}")
                return None
                
    if all_data:
        progress_bar.progress(1.0)
        status_text.success("🌿 **[ 100% / 100% ]** 모든 영단어 정제 프로세스가 성공적으로 완료되었습니다!")
    return all_data

# ==========================================
# 4. Streamlit 메인 UI 대시보드
# ==========================================
st.set_page_config(page_title="Voca-converter", layout="centered", page_icon="📝")

st.markdown("""
    <style>
    html, body, [data-testid="stAppViewContainer"] {
        background-color: #FBF9F4 !important;
    }
    
    [data-testid="stMainBlockContainer"] {
        background-color: transparent !important;
        max-width: 720px !important;
        margin: 0 auto !important;
        padding-top: 50px !important;
    }
    
    [data-testid="stVerticalBlockBorderContainer"] {
        border: none !important;
        background: transparent !important;
        box-shadow: none !important;
        padding: 0 !important;
    }

    .brand-title {
        font-size: 52px !important;
        font-weight: 700 !important;
        color: #556B2F !important;
        text-align: center !important;
        margin-bottom: 5px !important;
        letter-spacing: -1px !important;
    }
    
    .brand-caption {
        font-size: 15px !important;
        color: #8C9A86 !important;
        text-align: center !important;
        margin-bottom: 5px !important;
        font-weight: 500 !important;
    }
    
    .brand-author {
        font-size: 13px !important;
        color: #A0ABA2 !important;
        text-align: right !important;
        margin-bottom: 45px !important;
        font-weight: 500 !important;
        padding-right: 5px;
    }

    [data-testid="stFileUploader"] {
        border: none !important;
        background-color: #EEF1F6 !important;
        border-radius: 14px !important;
        padding: 20px 25px !important;
    }
    
    div.stButton > button:first-child {
        background-color: #85A392 !important; 
        color: white !important;
        border: none !important;
        font-size: 16px !important;
        font-weight: 500 !important;
        border-radius: 10px !important;
        padding: 12px 24px !important;
        box-shadow: none !important;
        width: auto !important;
    }
    div.stButton > button:first-child:hover {
        background-color: #6C8B7A !important;
    }
    
    [data-testid="stDownloadButton"]>button {
        background-color: #78909C !important;
        color: white !important;
        border-radius: 10px !important;
        border: none !important;
        padding: 12px 24px !important;
    }
    [data-testid="stDownloadButton"]>button:hover {
        background-color: #607D8B !important;
    }
    
    div[data-testid="stNotification"] {
        background-color: #E8F1FC !important;
        border: none !important;
        border-radius: 12px !important;
    }
    div[data-testid="stNotification"] p {
        color: #1E60B4 !important;
        font-weight: 500 !important;
    }
    
    .stProgress > div > div > div > div {
        background-color: #85A392 !important;
    }
    </style>
""", unsafe_allow_html=True)

st.markdown("<div class='brand-title'>Voca-converter</div>", unsafe_allow_html=True)
st.markdown("<div class='brand-caption'>사진 속 지문을 인식하여 편집 가능한 워드 문서(.docx)로 변환합니다.</div>", unsafe_allow_html=True)
st.markdown("<div class='brand-author'>(Made by Manju)</div>", unsafe_allow_html=True)

if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]
else:
    st.error("❌ Streamlit Cloud 설정의 Secrets에 GEMINI_API_KEY가 등록되지 않았습니다.")
    st.stop()

uploaded_files = st.file_uploader(
    "변환할 영어 지문 사진을 업로드하세요 (복수 선택 가능)", 
    type=["jpg", "jpeg", "png"], 
    accept_multiple_files=True
)

if uploaded_files:
    st.write("")
    st.markdown(f"📂 **{len(uploaded_files)}개의 파일이 선택되었습니다.**")
    
    if st.button("Word 파일로 변환하기 ✨", type="primary"):
        client = genai.Client(api_key=api_key)
        
        status_text = st.empty()
        progress_bar = st.progress(0)
        
        all_word_data = process_images_safely(client, uploaded_files, api_key, progress_bar, status_text)
        
        if all_word_data:
            st.toast("단어 데이터 정제가 완료되었습니다!")
            st.write("---")
            st.write("### 🔍 데이터 통합 미리보기")
            st.dataframe(all_word_data, use_container_width=True)
            
            word_file_buffer = create_word_document(all_word_data)
            
            st.write("")
            st.download_button(
                label="📥 정제된 수업용 Word 문서 다운로드 (.docx)",
                data=word_file_buffer,
                file_name="🔮_통합_영어단어장.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
